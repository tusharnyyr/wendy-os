"""
Create comprehensive Word documentation for Wendy v0.1
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def add_code_block(doc, code_text):
    """Add a code block with monospace font."""
    p = doc.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    return p

def create_wendy_documentation():
    """Create complete Wendy documentation in Word format."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Wendy v0.1', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Personal AI Automation Assistant')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.font.size = Pt(14)
    subtitle_format.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph(f'Documentation Generated: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('Version: 0.1.0 - Core Logging Engine')
    doc.add_paragraph('Status: ✅ Complete')
    
    doc.add_page_break()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Overview',
        '2. Features',
        '3. Architecture',
        '4. Installation Guide',
        '5. Project Structure',
        '6. Database Schema',
        '7. API Reference',
        '8. Usage Examples',
        '9. Development Steps',
        '10. Troubleshooting',
        '11. Cost Analysis',
        '12. Future Roadmap'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # 1. Overview
    doc.add_heading('1. Overview', 1)
    doc.add_paragraph(
        'Wendy is a modular AI automation framework that helps you track activities, '
        'build habits, and automate your personal workflows. Version 0.1 focuses on '
        'core activity logging with AI-powered parsing, weekly analytics, and streak tracking.'
    )
    
    doc.add_heading('Key Highlights', 2)
    doc.add_paragraph('Natural language input processing using Anthropic Claude', style='List Bullet')
    doc.add_paragraph('Multi-user support with flexible authentication', style='List Bullet')
    doc.add_paragraph('Automatic weekly analytics and streak detection', style='List Bullet')
    doc.add_paragraph('Event-driven architecture with webhook support', style='List Bullet')
    doc.add_paragraph('PostgreSQL database for reliable data persistence', style='List Bullet')
    
    # 2. Features
    doc.add_heading('2. Features', 1)
    
    doc.add_heading('✅ Implemented in v0.1', 2)
    doc.add_paragraph('Multi-user authentication via WhatsApp phone or API token', style='List Bullet')
    doc.add_paragraph('Natural language logging - "Log: Studied Python for 2 hours"', style='List Bullet')
    doc.add_paragraph('AI-powered parsing using Claude Haiku', style='List Bullet')
    doc.add_paragraph('Category classification (Learning, Work, Fitness, Personal, Project)', style='List Bullet')
    doc.add_paragraph('Smart duration parsing (2h, 45m, 1.5 hours)', style='List Bullet')
    doc.add_paragraph('Weekly total calculations (Monday to Sunday)', style='List Bullet')
    doc.add_paragraph('Streak tracking (consecutive days)', style='List Bullet')
    doc.add_paragraph('Event emission system with n8n webhook support', style='List Bullet')
    
    doc.add_heading('🚧 Planned for Future Versions', 2)
    doc.add_paragraph('Task management system', style='List Bullet')
    doc.add_paragraph('Goal tracking with milestones', style='List Bullet')
    doc.add_paragraph('Scheduled weekly summaries', style='List Bullet')
    doc.add_paragraph('AI-powered intent classification', style='List Bullet')
    doc.add_paragraph('Advanced analytics dashboard', style='List Bullet')
    doc.add_paragraph('WhatsApp/Telegram integration', style='List Bullet')
    
    # 3. Architecture
    doc.add_heading('3. Architecture', 1)
    
    doc.add_heading('Tech Stack', 2)
    tech_table = doc.add_table(rows=6, cols=2)
    tech_table.style = 'Light Grid Accent 1'
    
    tech_data = [
        ['Component', 'Technology'],
        ['Backend Framework', 'FastAPI (Python 3.10+)'],
        ['Database', 'PostgreSQL 18'],
        ['ORM', 'SQLAlchemy 2.0'],
        ['AI Provider', 'Anthropic Claude (Haiku model)'],
        ['Automation', 'n8n webhooks (optional)']
    ]
    
    for i, (component, tech) in enumerate(tech_data):
        tech_table.rows[i].cells[0].text = component
        tech_table.rows[i].cells[1].text = tech
    
    doc.add_paragraph()
    
    doc.add_heading('Request Flow', 2)
    doc.add_paragraph('1. User sends message (WhatsApp/Web/API)', style='List Number')
    doc.add_paragraph('2. Authentication layer resolves user', style='List Number')
    doc.add_paragraph('3. Intent router detects command type (Log: or Chat)', style='List Number')
    doc.add_paragraph('4. For logs: AI adapter parses natural language → JSON', style='List Number')
    doc.add_paragraph('5. Logging service saves to database', style='List Number')
    doc.add_paragraph('6. Analytics computed (weekly total, streak)', style='List Number')
    doc.add_paragraph('7. Event bus emits log_created event', style='List Number')
    doc.add_paragraph('8. Optional: Webhook sent to n8n', style='List Number')
    doc.add_paragraph('9. Formatted response returned to user', style='List Number')
    
    doc.add_page_break()
    
    # 4. Installation Guide
    doc.add_heading('4. Installation Guide', 1)
    
    doc.add_heading('Prerequisites', 2)
    doc.add_paragraph('Python 3.10 or higher', style='List Bullet')
    doc.add_paragraph('PostgreSQL 13+ (18 recommended)', style='List Bullet')
    doc.add_paragraph('Anthropic API key (free $5 credits available)', style='List Bullet')
    doc.add_paragraph('Windows 11 / macOS / Linux', style='List Bullet')
    
    doc.add_heading('Step-by-Step Setup', 2)
    
    doc.add_paragraph('Step 1: Create Project Directory').bold = True
    add_code_block(doc, 'cd Documents\nmkdir wendy\ncd wendy')
    
    doc.add_paragraph('Step 2: Create Virtual Environment').bold = True
    add_code_block(doc, 'python -m venv venv\nvenv\\Scripts\\activate  # Windows')
    
    doc.add_paragraph('Step 3: Install Dependencies').bold = True
    add_code_block(doc, 'pip install fastapi uvicorn sqlalchemy psycopg2-binary python-dotenv pydantic anthropic requests')
    
    doc.add_paragraph('Step 4: Configure Environment').bold = True
    add_code_block(doc, 'DATABASE_URL=postgresql://postgres:admin@localhost:5432/wendy\nANTHROPIC_API_KEY=sk-ant-your-key-here\nAI_PROVIDER=anthropic')
    
    doc.add_paragraph('Step 5: Create PostgreSQL Database').bold = True
    add_code_block(doc, 'psql -U postgres\nCREATE DATABASE wendy;\n\\q')
    
    doc.add_paragraph('Step 6: Initialize Database').bold = True
    add_code_block(doc, 'python init_database.py')
    
    doc.add_paragraph('Step 7: Run Application').bold = True
    add_code_block(doc, 'python main.py\nServer starts at: http://localhost:8000')
    
    doc.add_page_break()
    
    # 5. Project Structure
    doc.add_heading('5. Project Structure', 1)
    
    add_code_block(doc, '''wendy/
├── database/
│   ├── __init__.py
│   ├── base.py              # SQLAlchemy engine & session
│   └── models.py            # Database models
├── services/
│   ├── __init__.py
│   ├── auth_service.py      # User authentication
│   ├── intent_router.py     # Rule-based routing
│   ├── ai_adapter.py        # Claude API integration
│   ├── logging_service.py   # Activity logging & analytics
│   └── event_bus.py         # Event emission & webhooks
├── venv/                    # Virtual environment
├── .env                     # Environment variables
├── main.py                  # FastAPI application
└── README.md                # Project documentation''')
    
    doc.add_page_break()
    
    # 6. Database Schema
    doc.add_heading('6. Database Schema', 1)
    
    doc.add_heading('users table', 2)
    users_table = doc.add_table(rows=7, cols=3)
    users_table.style = 'Light Grid Accent 1'
    
    users_data = [
        ['Column', 'Type', 'Description'],
        ['id', 'UUID', 'Primary key'],
        ['name', 'String(255)', 'User name'],
        ['email', 'String(255)', 'Unique email'],
        ['whatsapp_number', 'String(20)', 'WhatsApp (unique, optional)'],
        ['api_token', 'String(255)', 'API token (unique, optional)'],
        ['created_at', 'DateTime', 'Creation timestamp']
    ]
    
    for i, (col, type_, desc) in enumerate(users_data):
        users_table.rows[i].cells[0].text = col
        users_table.rows[i].cells[1].text = type_
        users_table.rows[i].cells[2].text = desc
    
    doc.add_paragraph()
    
    doc.add_heading('activity_logs table', 2)
    logs_table = doc.add_table(rows=9, cols=3)
    logs_table.style = 'Light Grid Accent 1'
    
    logs_data = [
        ['Column', 'Type', 'Description'],
        ['id', 'UUID', 'Primary key'],
        ['user_id', 'UUID', 'Foreign key to users'],
        ['date', 'Date', 'Activity date'],
        ['activity_name', 'String(255)', 'Activity description'],
        ['category', 'Enum', 'Learning, Work, Fitness, Personal, Project'],
        ['duration_minutes', 'Integer', 'Duration in minutes'],
        ['notes', 'Text', 'Optional notes'],
        ['created_at', 'DateTime', 'Creation timestamp']
    ]
    
    for i, (col, type_, desc) in enumerate(logs_data):
        logs_table.rows[i].cells[0].text = col
        logs_table.rows[i].cells[1].text = type_
        logs_table.rows[i].cells[2].text = desc
    
    doc.add_page_break()
    
    # 7. API Reference
    doc.add_heading('7. API Reference', 1)
    
    doc.add_heading('POST /message', 2)
    doc.add_paragraph('Main endpoint for processing user messages.')
    doc.add_paragraph('Headers: Authorization: Bearer <token>')
    
    add_code_block(doc, '{\n  "message": "Log: Studied Python for 2 hours",\n  "phone_number": "+919876543210"\n}')
    
    doc.add_paragraph()
    
    doc.add_heading('POST /users/create', 2)
    doc.add_paragraph('Create a new user for testing.')
    doc.add_paragraph('Parameters: name, email, whatsapp_number (optional), api_token (optional)')
    
    doc.add_paragraph()
    
    doc.add_heading('GET /health', 2)
    doc.add_paragraph('Health check endpoint returns: {"status": "healthy", "service": "Wendy v0.1"}')
    
    doc.add_page_break()
    
    # 8. Usage Examples
    doc.add_heading('8. Usage Examples', 1)
    
    doc.add_heading('Create User', 2)
    add_code_block(doc, 'curl -X POST "http://localhost:8000/users/create?name=John&email=john@example.com&api_token=test123"')
    
    doc.add_paragraph()
    
    doc.add_heading('Log Activity', 2)
    add_code_block(doc, 'curl -X POST "http://localhost:8000/message" \\\n  -H "Authorization: Bearer test123" \\\n  -d \'{"message": "Log: Studied Python for 2 hours"}\'')
    
    doc.add_page_break()
    
    # 9-12. Other sections
    doc.add_heading('9. Development Steps Completed', 1)
    doc.add_paragraph('✅ Step 0: Prerequisites Setup')
    doc.add_paragraph('✅ Step 1: Database Models')
    doc.add_paragraph('✅ Step 2: Authentication Service')
    doc.add_paragraph('✅ Step 3: Intent Router')
    doc.add_paragraph('✅ Step 4: AI Adapter')
    doc.add_paragraph('✅ Step 5: Logging Service')
    doc.add_paragraph('✅ Step 6: Event Bus')
    doc.add_paragraph('✅ Step 7: FastAPI Application')
    
    doc.add_page_break()
    
    doc.add_heading('10. Cost Analysis', 1)
    doc.add_paragraph('Model: Claude Haiku')
    doc.add_paragraph('Cost per log: ~$0.0001 (0.01 cents)')
    doc.add_paragraph('Free $5 credits: 30,000-50,000 logs')
    doc.add_paragraph('Duration: 4-13 years for personal use')
    
    doc.add_page_break()
    
    doc.add_heading('11. Future Roadmap', 1)
    doc.add_paragraph('v0.2 - Task Management', style='List Bullet')
    doc.add_paragraph('v0.3 - Goal Tracking', style='List Bullet')
    doc.add_paragraph('v0.4 - Scheduled Summaries', style='List Bullet')
    doc.add_paragraph('v0.5 - WhatsApp Integration', style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('---')
    
    footer = doc.add_paragraph('Built with ❤️ using Claude (Anthropic)')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    doc.save('Wendy_v0.1_Complete_Documentation.docx')
    print("✅ Documentation created: Wendy_v0.1_Complete_Documentation.docx")

if __name__ == "__main__":
    create_wendy_documentation()